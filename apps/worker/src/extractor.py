"""Resume text extraction module supporting PDF, DOCX, and TXT files per Engineering Guidelines §6."""

import io

import docx
import pdfplumber
import structlog

from hiron.resumes.exceptions import ResumeParseFailedError

logger = structlog.get_logger("hiron.resumes.extractor")


MAX_RESUME_TEXT_CHARS = 30_000


def extract_text_from_pdf(file_bytes: bytes, max_chars: int) -> tuple[str, bool]:
    """Extract plain text from PDF file bytes using pdfplumber."""
    try:
        text_parts: list[str] = []
        total_chars = 0
        is_truncated = False
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    if total_chars + len(page_text) > max_chars:
                        remaining = max_chars - total_chars
                        text_parts.append(page_text[:remaining])
                        is_truncated = True
                        break

                    text_parts.append(page_text)
                    total_chars += len(page_text)
                else:
                    logger.debug("Empty or scanned page in PDF", page=page_idx)

        extracted = "\n\n".join(text_parts).strip()
        # Ensure exact limit after joining with newlines
        if len(extracted) > max_chars:
            extracted = extracted[:max_chars]
            is_truncated = True

        if not extracted:
            raise ResumeParseFailedError(
                "Unable to extract text from PDF (file may be empty, scanned, or encrypted)"
            )
        return extracted, is_truncated
    except Exception as exc:
        if isinstance(exc, ResumeParseFailedError):
            raise
        logger.warning("PDF text extraction failed", error=str(exc))
        raise ResumeParseFailedError(f"PDF extraction failed: {exc}") from exc


def _extract_docx_tables(
    tables: list, text_parts: list[str], total_chars: int, max_chars: int
) -> tuple[int, bool]:
    """Helper to extract text from DOCX tables while respecting the max_chars limit."""
    is_truncated = False
    for table in tables:
        if is_truncated:
            break
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                if total_chars + len(row_text) > max_chars:
                    remaining = max_chars - total_chars
                    text_parts.append(row_text[:remaining])
                    is_truncated = True
                    break
                text_parts.append(row_text)
                total_chars += len(row_text)
    return total_chars, is_truncated


def extract_text_from_docx(file_bytes: bytes, max_chars: int) -> tuple[str, bool]:
    """Extract plain text from DOCX file bytes using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts: list[str] = []
        total_chars = 0
        is_truncated = False

        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text = paragraph.text.strip()
                if total_chars + len(text) > max_chars:
                    remaining = max_chars - total_chars
                    text_parts.append(text[:remaining])
                    is_truncated = True
                    break
                text_parts.append(text)
                total_chars += len(text)

        if not is_truncated:
            total_chars, is_truncated = _extract_docx_tables(
                doc.tables, text_parts, total_chars, max_chars
            )

        extracted = "\n".join(text_parts).strip()
        if len(extracted) > max_chars:
            extracted = extracted[:max_chars]
            is_truncated = True

        if not extracted:
            raise ResumeParseFailedError(
                "Unable to extract text from DOCX (document appears empty)"
            )
        return extracted, is_truncated
    except Exception as exc:
        if isinstance(exc, ResumeParseFailedError):
            raise
        logger.warning("DOCX text extraction failed", error=str(exc))
        raise ResumeParseFailedError(f"DOCX extraction failed: {exc}") from exc


def extract_text_from_txt(file_bytes: bytes, max_chars: int) -> tuple[str, bool]:
    """Extract text from plain text file bytes with encoding fallback."""
    try:
        decoded = file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        try:
            decoded = file_bytes.decode("latin-1").strip()
        except Exception as exc:
            raise ResumeParseFailedError(f"TXT decoding failed: {exc}") from exc

    if len(decoded) > max_chars:
        return decoded[:max_chars], True
    return decoded, False


def extract_text_from_file(file_bytes: bytes, content_type: str, filename: str) -> tuple[str, bool]:
    """Dispatch text extraction based on MIME type or filename extension.
    Returns (extracted_text, is_truncated) tuple.
    """
    if not file_bytes:
        raise ResumeParseFailedError("File content is empty (0 bytes)")

    ext = "." + filename.split(".")[-1].lower() if "." in filename else ""

    if content_type == "application/pdf" or ext == ".pdf":
        return extract_text_from_pdf(file_bytes, MAX_RESUME_TEXT_CHARS)
    if (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or ext == ".docx"
    ):
        return extract_text_from_docx(file_bytes, MAX_RESUME_TEXT_CHARS)
    if content_type == "text/plain" or ext == ".txt":
        return extract_text_from_txt(file_bytes, MAX_RESUME_TEXT_CHARS)

    raise ResumeParseFailedError(
        f"Unsupported content type '{content_type}' or extension '{ext}' for text extraction"
    )
